"""
rag_system/skills/purchase_order.py

The first TASK-route skill: drafts a purchase order as a real .docx, grounded in
the structured warehouse the file router populated. localGPT's RAG path can only
answer questions; this *produces a deliverable*.

It is deliberately conservative: it pulls vendor/rate/line-items from DuckDB. If
it can't find the data it needs, it says so instead of fabricating a PO.

Requires: python-docx  (pip install python-docx)
"""
from datetime import date
import os
import re

import duckdb

DB_PATH = os.environ.get("STRUCTURED_DB", "./index_store/structured.duckdb")


def _inr(n):
    try:
        return "Rs " + format(int(round(float(n))), ",d")
    except (TypeError, ValueError):
        return str(n)


def _find_event(con, query):
    """Best-effort: pull an event id from the query, else the highest-value event."""
    m = re.search(r"\bEVN\s?\d+\b", query, re.I)
    tables = [t[0] for t in con.execute("SHOW TABLES").fetchall()]
    if "bids" not in tables or "events" not in tables:
        return None
    if m:
        eid = m.group(0).upper().replace(" ", "")
        row = con.execute("""
            SELECT b.event_id, b.l1_transporter vendor, b.item_desc, b.vehicle_qty qty, b.l1_rate rate,
                   e.origin, e.dest
            FROM bids b JOIN events e USING(event_id)
            WHERE b.is_winner AND b.event_id = ? AND b.l1_rate IS NOT NULL
            LIMIT 1""", [eid]).df()
        if len(row):
            return row.iloc[0]
    row = con.execute("""
        SELECT b.event_id, b.l1_transporter vendor, b.item_desc, b.vehicle_qty qty, b.l1_rate rate,
               e.origin, e.dest
        FROM bids b JOIN events e USING(event_id)
        WHERE b.is_winner AND b.l1_rate IS NOT NULL AND b.l1_transporter IS NOT NULL
        ORDER BY b.l1_rate DESC LIMIT 1""").df()
    return row.iloc[0] if len(row) else None


def draft_purchase_order(query, out_dir="./index_store/outputs", db_path=DB_PATH):
    """Return {'answer', 'task_file'} or {'answer'} if data is insufficient."""
    if not os.path.exists(db_path):
        return {"answer": "No structured data is indexed yet, so I can't draft a grounded "
                          "purchase order. Upload the sourcing/bid data first."}
    con = duckdb.connect(db_path)
    ev = _find_event(con, query)
    con.close()
    if ev is None:
        return {"answer": "I couldn't find awarded vendor/rate data in the indexed tables, "
                          "so I won't fabricate a purchase order. Please index the bid data "
                          "or specify an event id (e.g. EVN3440)."}

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"answer": "python-docx is not installed (pip install python-docx)."}

    qty = int(ev["qty"]) if ev["qty"] == ev["qty"] else 1
    amount = qty * float(ev["rate"])
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, f"purchase_order_{ev['event_id']}.docx")

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)
    h = doc.add_heading("Purchase Order", level=0)
    doc.add_paragraph(f"PO No: ESC/PO/{date.today():%Y}/{ev['event_id']}").runs[0].bold = True
    doc.add_paragraph(f"Date: {date.today():%Y-%m-%d}")
    doc.add_paragraph(f"Sourcing event: {ev['event_id']}    Lane: {ev['origin']} -> {ev['dest']}")
    doc.add_paragraph("Buyer", style="Heading 2")
    doc.add_paragraph("Escorts Kubota Limited, Faridabad, Haryana 121003, India")
    doc.add_paragraph("Supplier", style="Heading 2")
    doc.add_paragraph(str(ev["vendor"]))
    doc.add_paragraph("Order details", style="Heading 2")

    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"
    for c, t in zip(tbl.rows[0].cells, ["Description", "Qty", "Rate", "Amount"]):
        c.paragraphs[0].add_run(t).bold = True
    r = tbl.add_row().cells
    r[0].text = str(ev["item_desc"]); r[1].text = str(qty)
    r[2].text = _inr(ev["rate"]); r[3].text = _inr(amount)
    tot = tbl.add_row().cells
    tot[0].paragraphs[0].add_run("Total").bold = True
    tot[3].paragraphs[0].add_run(_inr(amount)).bold = True

    doc.add_paragraph("Terms", style="Heading 2")
    doc.add_paragraph("Net 30 days. Orders above Rs 1,00,000 require dual approval per vendor policy.")
    doc.add_paragraph("\nAuthorised signatory: ____________________    Date: __________")
    doc.save(out)

    return {"answer": f"Drafted a purchase order for {ev['vendor']} covering the "
                      f"{ev['origin']}->{ev['dest']} lane (event {ev['event_id']}) at {_inr(amount)}. "
                      f"File: {out}",
            "task_file": out}


# simple intent dispatcher for the task route
def handle_task(query, **kw):
    if re.search(r"purchase order|\bP\.?O\.?\b|raise an order|draft.*order", query, re.I):
        return draft_purchase_order(query, **kw)
    return {"answer": "No matching task skill. Available skills: purchase_order."}
